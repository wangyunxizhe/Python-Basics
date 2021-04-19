from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE

# 1，创建一个文档对象
document = Document()
# 2，写入内容
document.add_heading("简介", level=4)  # 标题，级别
# 样式：自定义的样式以键值对的形式存在，如 textstyle=自定义样式
style = document.styles.add_style("textstyle", WD_STYLE_TYPE.PARAGRAPH)
print(style.style_id)
print(style.name)
style.font.size = Pt(5)

# 删除自定义样式：删除后，在下文中引用该样式的地方将会报错
# document.styles['textstyle'].delete()

# 段落
p1 = document.add_paragraph("aaaaaaaaaaaaaaa,bbbbbbbbbbbbbb,"
                            "cccccccccccccccccccccccccccc,dddddddddddddddddddddd.", style="textstyle")
p1.insert_paragraph_before("!!在段落前插入一个新的段落")
format = p1.paragraph_format
format.left_indent = Pt(20)  # 左缩进
format.right_indent = Pt(20)  # 右缩进
format.first_line_indent = Pt(20)  # 首行缩进
format.line_spacing = 1.5  # 行间距
run = p1.add_run("eeeeeeeeeeeeeee,fffffffffffffffffff,ggggggggggggggggggg.")  # 追加一行，并设置样式
run.font.size = Pt(12)
run.font.name = "微软雅黑"
run.font.color.rgb = RGBColor(235, 33, 24)
run1 = p1.add_run("hhhhhhhhhhhhhhhh,iiiiiiiiiiiiiiiii,jjjjjjjjjjjjjjjjjjjj,kkkkkkkkkkk.")  # 再追加一行，并设置样式
run1.bold = True  # 加粗
run1.font.underline = True  # 下划线
run1.font.italic = True  # 斜体

# 插入图片，并设置图片大小
document.add_picture("../../resources/吴亦凡.jpg", Pt(50), Pt(60))

# 插入表格
table = document.add_table(rows=1, cols=3)
headerCells = table.rows[0].cells
headerCells[0].text = "月份"
headerCells[1].text = "预期销售额"
headerCells[2].text = "实际销售额"
# 数据
data = (["一月份", 500, 450], ["二月份", 600, 450], ["三月份", 700, 550])
for item in data:
    rowsCells = table.add_row().cells
    rowsCells[0].text = item[0]
    rowsCells[1].text = str(item[1])
    rowsCells[2].text = str(item[2])

# 获取word中的表格
mt = document.tables[0]  # 获取word中的第一个表格（也就是上面的代码刚刚创建在word中的表格）
print(len(mt.rows))  # 打印总行数
print(len(mt.columns))  # 打印总列数
print(document.tables[0].cell(0, 2).text)  # 打印指定单元格

# 3，保存文档
document.save("../../resources/简介.docx")
